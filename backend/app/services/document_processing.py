"""
Document Processing Service
Handles PDF and Word document parsing and key term extraction using LOCAL embeddings
No paid API required - uses SentenceTransformers (already in requirements.txt)
"""

from typing import Dict, Any, List, Optional
from loguru import logger
from collections import Counter
import io
import re

# Try to import sentence_transformers (already in requirements.txt)
try:
    from sentence_transformers import SentenceTransformer, util
    import numpy as np
    HAS_EMBEDDINGS = True
    logger.info("[DocumentProcessor] SentenceTransformers available for semantic keyword extraction")
except ImportError:
    HAS_EMBEDDINGS = False
    logger.warning("[DocumentProcessor] SentenceTransformers not available, using frequency fallback")


class DocumentProcessingService:
    """
    Service for processing uploaded documents (PDF, Word)
    
    Features:
    - Extract text from PDF files
    - Extract text from Word documents (.docx)
    - Clean and normalize extracted text
    - Extract key terms using LOCAL SentenceTransformers (FREE, no API needed)
    """
    
    # Singleton pattern for embedding model - load once, reuse
    _embedding_model = None
    
    @classmethod
    def get_embedding_model(cls):
        """Get or initialize the embedding model (singleton)"""
        if cls._embedding_model is None and HAS_EMBEDDINGS:
            try:
                # Use the same model as your embeddings.py for consistency
                # 'all-MiniLM-L6-v2' is small, fast, and good quality
                cls._embedding_model = SentenceTransformer('all-MiniLM-L6-v2')
                logger.info("[DocumentProcessor] Loaded embedding model: all-MiniLM-L6-v2")
            except Exception as e:
                logger.error(f"[DocumentProcessor] Failed to load embedding model: {e}")
        return cls._embedding_model
    
    def __init__(self, llm_client=None):
        """
        Initialize the document processing service
        
        Args:
            llm_client: Optional LLM client (not used - we use local embeddings now)
        """
        # LLM client kept for backward compatibility but not used
        self.llm_client = llm_client
    
    async def process_document(
        self,
        file_content: bytes,
        filename: str,
        content_type: str
    ) -> Dict[str, Any]:
        """
        Process an uploaded document and extract text + key terms
        
        Args:
            file_content: Raw file bytes
            filename: Original filename
            content_type: MIME type of the file
            
        Returns:
            Dictionary with extracted_text, key_terms, and metadata
        """
        try:
            # Determine file type and extract text
            if content_type == "application/pdf" or filename.lower().endswith(".pdf"):
                extracted_text = await self._extract_pdf_text(file_content)
            elif content_type in [
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                "application/msword"
            ] or filename.lower().endswith((".docx", ".doc")):
                extracted_text = await self._extract_word_text(file_content, filename)
            elif content_type == "text/plain" or filename.lower().endswith(".txt"):
                extracted_text = file_content.decode("utf-8", errors="ignore")
            else:
                raise ValueError(f"Unsupported file type: {content_type}")
            
            # Clean the extracted text
            cleaned_text = self._clean_text(extracted_text)
            
            # Extract key terms using LOCAL embeddings (FREE)
            key_terms = self._extract_key_terms_local(cleaned_text)
            
            # Generate topic info from content analysis
            topic_info = self._generate_topic_info_local(cleaned_text, key_terms)
            
            return {
                "status": "success",
                "extracted_text": cleaned_text,
                "key_terms": key_terms,
                "topic_info": topic_info,
                "metadata": {
                    "filename": filename,
                    "content_type": content_type,
                    "text_length": len(cleaned_text),
                    "term_count": len(key_terms)
                }
            }
            
        except Exception as e:
            logger.error(f"Document processing failed: {e}")
            return {
                "status": "error",
                "error": str(e),
                "extracted_text": "",
                "key_terms": [],
                "topic_info": {},
                "metadata": {"filename": filename}
            }
    
    async def _extract_pdf_text(self, file_content: bytes) -> str:
        """Extract text from PDF using pypdf"""
        try:
            from pypdf import PdfReader
            
            pdf_file = io.BytesIO(file_content)
            reader = PdfReader(pdf_file)
            
            text_parts = []
            for page in reader.pages:
                page_text = page.extract_text()
                if page_text:
                    text_parts.append(page_text)
            
            return "\n\n".join(text_parts)
            
        except ImportError:
            logger.error("pypdf not installed. Install with: pip install pypdf")
            raise ValueError("PDF processing not available. pypdf package required.")
        except Exception as e:
            logger.error(f"PDF extraction error: {e}")
            raise ValueError(f"Failed to extract text from PDF: {e}")
    
    async def _extract_word_text(self, file_content: bytes, filename: str) -> str:
        """Extract text from Word documents"""
        try:
            from docx import Document
            
            doc_file = io.BytesIO(file_content)
            doc = Document(doc_file)
            
            text_parts = []
            
            # Extract text from paragraphs
            for para in doc.paragraphs:
                if para.text.strip():
                    text_parts.append(para.text)
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                    if row_text:
                        text_parts.append(row_text)
            
            return "\n\n".join(text_parts)
            
        except ImportError:
            logger.error("python-docx not installed. Install with: pip install python-docx")
            raise ValueError("Word document processing not available. python-docx package required.")
        except Exception as e:
            logger.error(f"Word extraction error: {e}")
            raise ValueError(f"Failed to extract text from Word document: {e}")
    
    def _clean_text(self, text: str) -> str:
        """Clean and normalize extracted text"""
        if not text:
            return ""
        
        # Remove excessive whitespace
        text = re.sub(r'\s+', ' ', text)
        
        # Remove special characters that might cause issues
        text = re.sub(r'[\x00-\x08\x0b\x0c\x0e-\x1f\x7f-\x9f]', '', text)
        
        # Normalize line breaks
        text = re.sub(r'\n{3,}', '\n\n', text)
        
        # Trim leading/trailing whitespace
        text = text.strip()
        
        # Truncate if too long (keep first 50k chars for processing)
        max_length = 50000
        if len(text) > max_length:
            text = text[:max_length]
            logger.warning(f"Document text truncated to {max_length} characters")
        
        return text
    
    def _extract_key_terms_local(self, text: str, top_k: int = 8) -> List[str]:
        """
        Extract key search terms using LOCAL SentenceTransformers embeddings.
        No paid API required - runs entirely on your server.
        
        Uses semantic similarity to find terms most representative of the document.
        Falls back to frequency-based extraction if embeddings unavailable.
        """
        if not text:
            return []
        
        # Limit processing to first 5000 chars for performance
        text_chunk = text[:5000]
        
        # Extract candidate terms (noun phrases, technical terms, etc.)
        candidates = self._extract_candidates(text_chunk)
        
        if not candidates:
            # Fallback: return first meaningful sentence fragment
            first_sentence = text_chunk.split('.')[0][:100].strip()
            return [first_sentence] if first_sentence else []
        
        # Try semantic ranking with embeddings
        if HAS_EMBEDDINGS:
            try:
                model = self.get_embedding_model()
                if model:
                    # Encode the full document chunk
                    doc_embedding = model.encode(text_chunk, convert_to_tensor=True)
                    
                    # Encode all candidates
                    candidate_embeddings = model.encode(candidates, convert_to_tensor=True)
                    
                    # Calculate cosine similarity between doc and each candidate
                    cos_scores = util.cos_sim(doc_embedding, candidate_embeddings)[0]
                    
                    # Get indices of top_k most similar candidates
                    scores_np = cos_scores.cpu().numpy()
                    top_indices = np.argsort(-scores_np)[:top_k]
                    
                    selected_terms = [candidates[i] for i in top_indices]
                    logger.info(f"[DocumentProcessor] Extracted {len(selected_terms)} key terms via semantic similarity")
                    return selected_terms
                    
            except Exception as e:
                logger.warning(f"[DocumentProcessor] Semantic extraction failed: {e}, using frequency fallback")
        
        # Fallback: frequency-based extraction
        return self._frequency_based_extraction(candidates, top_k)
    
    def _extract_candidates(self, text: str) -> List[str]:
        """
        Extract potential keyword candidates using regex heuristics.
        Targets technical terms, proper nouns, and meaningful phrases.
        """
        candidates = []
        
        # Extended stopwords list
        stopwords = {
            'the', 'is', 'in', 'at', 'of', 'on', 'and', 'a', 'an', 'to', 'for', 
            'with', 'by', 'this', 'that', 'it', 'are', 'was', 'were', 'from',
            'be', 'been', 'being', 'have', 'has', 'had', 'do', 'does', 'did',
            'will', 'would', 'could', 'should', 'may', 'might', 'must', 'shall',
            'can', 'need', 'dare', 'ought', 'used', 'into', 'through', 'during',
            'before', 'after', 'above', 'below', 'between', 'under', 'again',
            'further', 'then', 'once', 'here', 'there', 'when', 'where', 'why',
            'how', 'all', 'each', 'few', 'more', 'most', 'other', 'some', 'such',
            'no', 'nor', 'not', 'only', 'own', 'same', 'so', 'than', 'too', 'very',
            'just', 'but', 'if', 'or', 'because', 'until', 'while', 'what', 'which',
            'who', 'whom', 'whose', 'also', 'however', 'therefore', 'thus', 'hence',
            'example', 'examples', 'figure', 'table', 'chapter', 'section', 'page',
            'see', 'following', 'these', 'those', 'they', 'them', 'their', 'its',
            'about', 'over', 'such', 'your', 'our', 'his', 'her', 'my', 'any',
        }
        
        # Strategy 1: Capitalized phrases (proper nouns, technical terms)
        # Matches "Machine Learning", "Data Science", "Python Programming"
        capitalized_phrases = re.findall(r'\b[A-Z][a-z]+(?:\s+[A-Z][a-z]+)+\b', text)
        candidates.extend([p for p in capitalized_phrases if p.lower() not in stopwords])
        
        # Strategy 2: Single capitalized words (not at sentence start)
        # Look for capitalized words that aren't after a period
        single_caps = re.findall(r'(?<![.!?]\s)[A-Z][a-z]{2,}', text)
        candidates.extend([w for w in single_caps if w.lower() not in stopwords and len(w) > 3])
        
        # Strategy 3: Technical terms (often lowercase but meaningful)
        # Words with 5+ characters that appear multiple times
        words = re.findall(r'\b[a-zA-Z]{5,}\b', text.lower())
        word_freq = Counter(words)
        frequent_words = [w for w, count in word_freq.items() 
                        if count >= 2 and w not in stopwords]
        candidates.extend(frequent_words[:20])  # Top 20 frequent words
        
        # Strategy 4: Hyphenated or compound terms (often technical)
        # Matches "object-oriented", "machine-learning"
        compound_terms = re.findall(r'\b[a-zA-Z]+-[a-zA-Z]+\b', text)
        candidates.extend(compound_terms)
        
        # Strategy 5: Acronyms (all caps, 2-6 letters)
        acronyms = re.findall(r'\b[A-Z]{2,6}\b', text)
        candidates.extend([a for a in acronyms if a not in {'THE', 'AND', 'FOR', 'NOT'}])
        
        # Deduplicate while preserving order
        seen = set()
        unique_candidates = []
        for c in candidates:
            c_lower = c.lower()
            if c_lower not in seen and len(c) > 2:
                seen.add(c_lower)
                unique_candidates.append(c)
        
        return unique_candidates
    
    def _frequency_based_extraction(self, candidates: List[str], top_k: int) -> List[str]:
        """
        Fallback: extract terms based on frequency.
        Used when embeddings are not available.
        """
        if not candidates:
            return []
        
        # Count occurrences
        counter = Counter([c.lower() for c in candidates])
        
        # Get original casing for top terms
        term_map = {c.lower(): c for c in candidates}
        top_terms = [term_map.get(term, term) for term, _ in counter.most_common(top_k)]
        
        logger.info(f"[DocumentProcessor] Extracted {len(top_terms)} key terms via frequency analysis")
        return top_terms
    
    def _generate_topic_info_local(
        self,
        text: str,
        key_terms: List[str]
    ) -> Dict[str, Any]:
        """
        Generate topic information from the document locally.
        Uses heuristics to determine subject area and difficulty.
        """
        if not text:
            return {
                "suggested_topic": "Uploaded Document",
                "subject_area": "General",
                "difficulty_hint": "medium"
            }
        
        # Generate suggested topic from key terms
        if key_terms:
            # Use top 2-3 key terms as topic
            suggested_topic = " ".join(key_terms[:3])
        else:
            # Use first meaningful phrase from text
            first_line = text.split('\n')[0][:50].strip()
            suggested_topic = first_line if first_line else "Uploaded Document"
        
        # Detect subject area based on keyword patterns
        subject_area = self._detect_subject_area(text, key_terms)
        
        # Estimate difficulty based on text complexity
        difficulty_hint = self._estimate_difficulty(text)
        
        return {
            "suggested_topic": suggested_topic,
            "subject_area": subject_area,
            "difficulty_hint": difficulty_hint
        }
    
    def _detect_subject_area(self, text: str, key_terms: List[str]) -> str:
        """
        Detect the subject area based on keyword patterns.
        """
        text_lower = text.lower()
        terms_lower = " ".join([t.lower() for t in key_terms])
        combined = text_lower + " " + terms_lower
        
        # Subject area detection patterns
        subject_patterns = {
            "Computer Science": [
                'algorithm', 'programming', 'software', 'database', 'python', 
                'java', 'code', 'function', 'variable', 'class', 'object',
                'api', 'web', 'server', 'network', 'data structure', 'compiler',
                'machine learning', 'artificial intelligence', 'neural network'
            ],
            "Mathematics": [
                'equation', 'theorem', 'proof', 'calculus', 'algebra', 'geometry',
                'matrix', 'vector', 'integral', 'derivative', 'probability',
                'statistics', 'mathematical', 'formula', 'coefficient'
            ],
            "Physics": [
                'force', 'energy', 'momentum', 'velocity', 'acceleration',
                'quantum', 'relativity', 'thermodynamics', 'electromagnetic',
                'particle', 'wave', 'frequency', 'wavelength', 'physics'
            ],
            "Chemistry": [
                'molecule', 'atom', 'element', 'compound', 'reaction', 'bond',
                'organic', 'inorganic', 'chemical', 'solution', 'acid', 'base',
                'oxidation', 'catalyst', 'periodic table'
            ],
            "Biology": [
                'cell', 'dna', 'gene', 'protein', 'organism', 'species',
                'evolution', 'ecology', 'anatomy', 'physiology', 'bacteria',
                'virus', 'enzyme', 'metabolism', 'photosynthesis'
            ],
            "Business": [
                'market', 'finance', 'investment', 'strategy', 'management',
                'revenue', 'profit', 'business', 'company', 'stock', 'economy',
                'trade', 'marketing', 'sales', 'customer'
            ],
            "Engineering": [
                'design', 'system', 'circuit', 'mechanical', 'electrical',
                'structural', 'materials', 'manufacturing', 'engineering',
                'specifications', 'prototype', 'testing'
            ],
            "Medicine": [
                'patient', 'diagnosis', 'treatment', 'symptom', 'disease',
                'clinical', 'medical', 'therapy', 'surgery', 'drug',
                'pharmaceutical', 'healthcare', 'hospital'
            ],
        }
        
        # Count matches for each subject
        subject_scores = {}
        for subject, patterns in subject_patterns.items():
            score = sum(1 for p in patterns if p in combined)
            if score > 0:
                subject_scores[subject] = score
        
        # Return subject with highest score, or "General" if no matches
        if subject_scores:
            return max(subject_scores, key=subject_scores.get)
        return "General"
    
    def _estimate_difficulty(self, text: str) -> str:
        """
        Estimate difficulty level based on text complexity.
        """
        if not text:
            return "medium"
        
        # Sample first 2000 characters
        sample = text[:2000]
        words = sample.split()
        
        if not words:
            return "medium"
        
        # Calculate average word length
        avg_word_length = sum(len(w) for w in words) / len(words)
        
        # Count complex indicators
        complex_indicators = [
            'advanced', 'complex', 'sophisticated', 'comprehensive', 'theoretical',
            'doctoral', 'research', 'analysis', 'methodology', 'hypothesis',
            'furthermore', 'nevertheless', 'notwithstanding', 'consequently'
        ]
        
        simple_indicators = [
            'basic', 'introduction', 'beginner', 'fundamental', 'simple',
            'overview', 'getting started', 'learn', 'easy', 'first step'
        ]
        
        text_lower = text.lower()
        complex_count = sum(1 for i in complex_indicators if i in text_lower)
        simple_count = sum(1 for i in simple_indicators if i in text_lower)
        
        # Determine difficulty
        if simple_count > complex_count and avg_word_length < 5.5:
            return "easy"
        elif complex_count > simple_count + 2 or avg_word_length > 6.5:
            return "hard"
        elif complex_count > simple_count + 4 or avg_word_length > 7:
            return "expert"
        else:
            return "medium"
